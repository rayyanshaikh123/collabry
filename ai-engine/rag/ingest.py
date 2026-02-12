"""
Document Ingestion Pipeline.

Handles document upload, text extraction, chunking, and vector storage.
"""

import os
import io
from typing import List, Dict, Any, Optional
from pathlib import Path
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from rag.vectorstore import add_documents


async def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """
    Extract text from uploaded file.
    
    Supports: PDF, DOCX, TXT, MD
    
    Args:
        file_content: Raw file bytes
        filename: Original filename with extension
    
    Returns:
        Extracted text content
    """
    file_ext = Path(filename).suffix.lower()
    
    if file_ext == ".txt" or file_ext == ".md":
        # Plain text
        return file_content.decode("utf-8", errors="ignore")
    
    elif file_ext == ".pdf":
        # PDF extraction
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf not installed. Install with: pip install pypdf")
        
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        
        return text
    
    elif file_ext == ".docx":
        # DOCX extraction
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")


def chunk_text(
    text: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> List[str]:
    """
    Split text into chunks using recursive character splitting.
    
    Args:
        text: Input text
        chunk_size: Maximum chunk size in characters
        chunk_overlap: Overlap between chunks
    
    Returns:
        List of text chunks
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    return splitter.split_text(text)


async def ingest_document(
    file_content: bytes,
    filename: str,
    user_id: str,
    notebook_id: str,
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> Dict[str, Any]:
    """
    Ingest a document: extract text, chunk, embed, and store.
    
    Args:
        file_content: Raw file bytes
        filename: Original filename
        user_id: User identifier
        notebook_id: Notebook identifier
        chunk_size: Chunk size for splitting
        chunk_overlap: Overlap between chunks
    
    Returns:
        Ingestion summary: {success, filename, chunks, notebook_id}
    """
    try:
        # 1. Extract text
        text = await extract_text_from_file(file_content, filename)
        
        if not text.strip():
            raise ValueError("No text extracted from file")
        
        # 2. Chunk text
        chunks = chunk_text(text, chunk_size, chunk_overlap)
        
        # 3. Create LangChain Documents
        # Explicitly include source_id in metadata for precise filtering
        source_id = (metadata or {}).get("source_id") or filename
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": filename,
                    "source_id": source_id,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        # 4. Add to vector store (embeddings generated internally)
        doc_ids = add_documents(documents, user_id, notebook_id)
        
        return {
            "success": True,
            "filename": filename,
            "chunks": len(chunks),
            "notebook_id": notebook_id,
            "document_ids": doc_ids,
            "text_length": len(text)
        }
    
    except Exception as e:
        return {
            "success": False,
            "filename": filename,
            "error": str(e)
        }


async def ingest_text(
    text: str,
    user_id: str,
    notebook_id: str,
    source_name: str = "direct_input",
    chunk_size: int = 1000,
    chunk_overlap: int = 200
) -> Dict[str, Any]:
    """
    Ingest raw text directly (no file upload).
    
    Args:
        text: Raw text content
        user_id: User identifier
        notebook_id: Notebook identifier
        source_name: Source identifier
        chunk_size: Chunk size for splitting
        chunk_overlap: Overlap between chunks
    
    Returns:
        Ingestion summary
    """
    try:
        # 1. Chunk text
        chunks = chunk_text(text, chunk_size, chunk_overlap)
        
        # 2. Create LangChain Documents
        # Use source_name as source_id for direct text input
        documents = [
            Document(
                page_content=chunk,
                metadata={
                    "source": source_name,
                    "source_id": source_name,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        # 3. Add to vector store
        doc_ids = add_documents(documents, user_id, notebook_id)
        
        return {
            "success": True,
            "source": source_name,
            "chunks": len(chunks),
            "notebook_id": notebook_id,
            "document_ids": doc_ids,
            "text_length": len(text)
        }
    
    except Exception as e:
        return {
            "success": False,
            "source": source_name,
            "error": str(e)
        }


async def batch_ingest(
    files: List[tuple[bytes, str]],
    user_id: str,
    notebook_id: str
) -> Dict[str, Any]:
    """
    Ingest multiple documents in batch.
    
    Args:
        files: List of (file_content, filename) tuples
        user_id: User identifier
        notebook_id: Notebook identifier
    
    Returns:
        Batch ingestion summary
    """
    results = []
    
    for file_content, filename in files:
        result = await ingest_document(
            file_content=file_content,
            filename=filename,
            user_id=user_id,
            notebook_id=notebook_id
        )
        results.append(result)
    
    success_count = sum(1 for r in results if r["success"])
    total_chunks = sum(r.get("chunks", 0) for r in results if r["success"])
    
    return {
        "success": success_count > 0,
        "total_files": len(files),
        "successful": success_count,
        "failed": len(files) - success_count,
        "total_chunks": total_chunks,
        "results": results
    }
